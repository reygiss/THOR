import xlrd
from docxtpl import DocxTemplate
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
from tkinter import END


###################################################################
# Suppression d'une table dans le document word
# INPUT : la table à supprimer


def delete_table(old_table):
    parent = old_table._element.getparent()  # on recupere le parent de la table
    parent.remove(old_table)  # suppression de la table
    old_table._element = None  # libération mémoire


###################################################################
# Suppression d'un paragraph dans le document word
# INPUT : Le paragraph a supprimé


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


###################################################################
# Indentation d'une table dans la page
# INPUT : table : La table objet de l'indentation
# INPUT : indent : la valeur numérique de l'indentation


def indent_table(table, indent):
    # noinspection PyProtectedMember
    tbl_pr = table._element.xpath('w:tblPr')  # récupération de l'élément XML de la table
    if tbl_pr:
        e = OxmlElement('w:tblInd')  # création d'un element XML pour l'indentation
        e.set(qn('w:w'), str(indent))  # set de la valeur e l'indentation
        e.set(qn('w:type'), 'dxa')  # type de l'indentation
        tbl_pr[0].append(e)  # ajout de l'elment XML à celui de la table


###################################################################
# Cette fonction permet d'appliquer une couleur de fond à une cellule sous word
# INPUT: cell : Cellule du tableau
# INPUT : shade : valeur hexa de la couleur de fond = #FFFFFF


def set_shade_cell(cell, shade):
    tcpr = cell._tc.get_or_add_tcPr()  # récupération de lélément XML correspondant à la cellule
    tcvalign = OxmlElement("w:shd")  # création d'un élément XML pour le background
    tcvalign.set(qn("w:fill"), shade)  # set de la valeur de la vcouleur de fond
    tcpr.append(tcvalign)  # ajout de l'élement XML à la cellule


###################################################################
# Cette fonction permet de lire une couleur de fond à une cellule sous word
# INPUT: cell : Cellule du tableau
# OUTPUT : valeur hexa de la couleur de fond = #FFFFFF


def get_shade_cell(cell):
    tcpr = cell._tc.get_or_add_tcpr()  # récupération de lélément XML correspondant à la cellule
    tcshd = tcpr.xpath("w:shd")  # création d'un élément XML pour le background
    if len(tcshd) > 0:
        return tcshd[0].get(qn("w:fill"))  # get de la valeur de la couleur de fond
    else:
        return ""


###################################################################
# Suppression d'une ligne de tableau word
# INPUT : table : table du document word
# INPUT : row : la ligne de la table


def remove_row(table, row):
    tbl = table._tbl  # récupération de l'élément table
    tr = row._tr  # récupération de la ligne
    tbl.remove(tr)  # suppression de la ligne dans la table
    row = tr = None  # libération mémoire


###################################################################
# chargement d'une echelle fixe a partir d'un fichier excel
# INPUT : excel : chemin du fichier excel
# INPUT : sheet : nom de la feuille excel
# INPUT : nbEnetete : nombre de lignes d'entete du tableau excel
# INPUT : log : pointeur sur a fenetre de journalisation


def load_echelle_fixe(excel, sheet, nbentete, log, thor):
    echelle = []
    try:
        wb = xlrd.open_workbook(excel, formatting_info=True)  # ouverture du fichier Excel
        sheet = wb.sheet_by_name(sheet)  # Récupération de la feuille dans le fichie Excel
    except:
        if thor["debug"]:  # si mode debug activ" # si mode debug activ"
            sys.exc_info()[0]  # On affiche l'erreur
            raise  # levée de l'erreur
        else:
            log.insert(END, "\nERROR : Erreur à l'ouverture du fichier Excel " + excel)
            return echelle
    try:
        # récupération de la taille du tableau, les fonctions max_row
        # et max_col retourne des valeurs éronnées
        maxrow = sheet.nrows  # Nombre de lignes dans le fichier Excel
        for x in range(nbentete, maxrow):
            rgb = sheet.cell_value(x, 3).split(',')
            echelle.append(EchelleFixe(sheet.cell_value(x, 2),
                                       "#" + str(RGBColor(int(rgb[0]),
                                                          int(rgb[1]),
                                                          int(rgb[2])))))
        return echelle
    except:
        if thor["debug"]:  # si mode debug activ" # si mode debug activ"
            sys.exc_info()[0]  # On affiche l'erreur
            raise  # levée de l'erreur
        else:
            log.insert(END, "\nERROR : Erreur à l'import de la legende " + excel)
            return echelle


###################################################################
# chargement d'une echelle calculée a partir d'un fichier excel
# INPUT : excel : chemin du fichier excel
# INPUT : sheet : nom de la feuille excel
# INPUT : nbentete : nombre de lignes d'entete du tableau excel
# INPUT : log : pointeur sur a fenetre de journalisation


def load_echelle_calculee(excel, sheet, nbentete, log, thor):
    echelle = []
    try:
        # ouverture du fichier Excel
        wb = xlrd.open_workbook(excel, formatting_info=True)
        # Récupération de la feuille dans le fichie Excel
        sheet = wb.sheet_by_name(sheet)
    except:
        if thor["debug"]:  # si mode debug activ" # si mode debug activ"
            sys.exc_info()[0]  # On affiche l'erreur
            raise  # levée de l'erreur
        else:
            log.insert(END, "\nERROR : Erreur à l'ouverture du fichier Excel " + excel)
            return echelle
    try:
        # récupération de la taille du tableau, les fonctions max_row et
        # max_col retourne des valeurs éronnées
        maxrow = sheet.nrows  # Nombre de lignes dans le fichier Excel
        for x in range(nbentete, maxrow):
            rgb = sheet.cell_value(x, 3).split(',')
            echelle.append(EchelleCalculee(sheet.cell_value(x, 1),
                                           sheet.cell_value(x, 2),
                                           "#" + str(RGBColor(int(rgb[0]),
                                                              int(rgb[1]),
                                                              int(rgb[2])))))
        return echelle
    except:
        if thor["debug"]:  # si mode debug activ" # si mode debug activ"
            sys.exc_info()[0]  # On affiche l'erreur
            raise  # levée de l'erreur
        else:
            log.insert(END, "\nERROR : Erreur \
            à l'import de la legende " + excel)


###################################################################
# cette fonction permet de copier un tableau de la feuille
# sheet du fichier excel dans le document word doc
# INPUT : doc : document word
# INPUT : index : index de la table dans le document word
# INPUT : nbenteteword : nombre de ligne dans l'entete de la table dans word
# INPUT : nbenteteexcel : nombre de ligne dans l'entete de la table dans excel
# INPUT : excel : Nom du fichier Excel
# INPUT : sheet : nom de la feuille dans Excel
# INPUT : nbColonnesIgnorees: nombre de colonnes à gauche qu'il faut ignorer, 0 par défaut
# INPUT : log : pointeur sur a fenetre de journalisation


def copy_table(doc, index, tab, thor, excel, log):
    nbenteteword = tab["enteteWord"]
    nbenteteexcel = tab["enteteExcel"]
    sheet = tab["feuilleExcel"]
    nbcolonnesignorees = tab["nbColonnesIgnorees"]
    textstyle = tab['style']['textStyle']
    table = doc.tables[index]  # Récupération de la table dans le document Word
    # Get Excel
    try:
        wb = xlrd.open_workbook(excel, formatting_info=True)  # ouverture du fichier Excel
        sheet = wb.sheet_by_name(sheet)  # Récupération de la feuille dans le fichier Excel
    except:
        if thor["debug"]:  # si mode debug activ" # si mode debug activ"
            sys.exc_info()[0]  # On affiche l'erreur
            raise  # levée de l'erreur
        else:
            log.insert(END, "\nERROR : Erreur à \
            l'ouverture du fichier Excel " + excel)
            return 1
    try:
        # récupération de la taille du tableau,
        # les fonctions max_ro  # w et max_col retourne des valeurs éronnées
        maxrow = sheet.nrows  # Nombre de lignes dans le fichier Excel
        maxcol = sheet.ncols  # Nombre de colonnes dans le fichier Excel

        # dans word on supprime toutes les lignes sauf l'entete
        # tant que le nombre de lignes est
        # superieur au nombre de ligne de l'entete
        while len(table.rows) > nbenteteword:
            # suppression de la derniere ligne
            remove_row(table, table.rows[nbenteteword])

        # Remplissage du tableau word avec les valeurs de l'Excel
        # Calcul de l'écart entre l'entete word et Excel
        # pour calculer les difference de numéro de ligne
        ecart = nbenteteword - nbenteteexcel
        # Pour chaque ligne du fichier Excel, sauf entete
        for x in range(nbenteteexcel, maxrow):
            r = table.add_row()  # on ajoute une ligne dans le tableau word
            # La ligne ne peut pas etre splitter sur plusieurs pages
            r._element.append(OxmlElement('w:cantSplit'))
            # On ignore la premiere colonne Excel qui
            # contient les numéro de ligne, sinon pour chaque colonne,
            # donc il y a une différence de 1 entre le
            # numéro de la colonne word et celle Excel
            for y in range(nbcolonnesignorees, maxcol):
                # type 0 = cellule vide, si la cellule n'est pas vide
                if sheet.cell_type(x, y) > 0:
                    # On recopie la valeur de la cellule
                    text_cell = str(sheet.cell_value(x, y))
                    # On recopie les styles de la cellule
                    text_cell_xf = wb.xf_list[sheet.cell_xf_index(x, y)]
                    # On recopie les multi-styles de la cellule
                    text_cell_runlist = sheet.rich_text_runlist_map.get((x, y))
                    if text_cell_runlist:  # Si styles multiples
                        segments = []
                        for segment_idx in range(len(text_cell_runlist)):  # liste des segments
                            start = text_cell_runlist[segment_idx][0]  # debut du segmment
                            # the last segment starts at given 'start'
                            # and ends at the end of the string
                            end = None
                            if segment_idx != len(text_cell_runlist) - 1:  # si pas dernier segment
                                end = text_cell_runlist[segment_idx + 1][0]
                                segment_text = text_cell[start:end]  # text du segment
                                segments.append({  # On ajoute le segment
                                    'text': segment_text,  # Text du segement
                                    'font': wb.font_list[text_cell_runlist[segment_idx][1]]  # style su segment
                                })
                        # segments did not start at beginning, assume cell
                        # starts with text styled as the cell
                        if text_cell_runlist[0][0] != 0:
                            segments.insert(0, {
                                'text': text_cell[:text_cell_runlist[0][0]],
                                'font': wb.font_list[text_cell_xf.font_index]
                            })

                        # On récupère le paragraph sous Word
                        p = table.cell(x + ecart, y - nbcolonnesignorees).paragraphs[0]
                        for segment in segments:  # Pour chaque segment
                            r = p.add_run()  # On ajoute un run à word
                            # On colle la valeur dans la cellule
                            # correspondante sous Word
                            r.text = segment['text']
                            # Si le segment est en italique
                            if segment['font'].italic:
                                r.font.italic = True
                            # Si le segment est en Gras
                            if segment['font'].bold:
                                r.font.bold = True
                            # Si le segment est souligné
                            if segment['font'].underlined:
                                r.font.underline = True
                            # Si le segment à une couleur de police
                            if segment['font'].colour_index:
                                color = wb.colour_map.get(segment['font'].colour_index)
                                if color:
                                    r.font.color.rgb = RGBColor(color[0], color[1], color[2])
                    else:  # La cellule n'a qu'un style simple
                        p = table.cell(x + ecart, y - nbcolonnesignorees).paragraphs[0]
                        r = p.add_run()
                        r.text = text_cell
                        # si italique
                        if wb.font_list[text_cell_xf.font_index].italic:
                            r.font.italic = True
                        # si souligne
                        if wb.font_list[text_cell_xf.font_index].underlined:
                            r.font.underline = True
                        if wb.font_list[text_cell_xf.font_index].bold:  # Si Gras
                            r.font.bold = True
                        if wb.font_list[text_cell_xf.font_index].colour_index:  # Si couleur de police
                            color = wb.colour_map.get(
                                wb.font_list[text_cell_xf.font_index].colour_index)
                            if color:
                                r.font.color.rgb = RGBColor(color[0], color[1], color[2])
                    # conversion des retour à la ligne en paragraphs
                    # Pour chaque paragraphs de la cellule
                    for p in table.cell(x + ecart, y - nbcolonnesignorees).paragraphs:
                        p2 = table.cell(x + ecart, y - nbcolonnesignorees).add_paragraph()
                        for r in p.runs:
                            lignes = r.text.split('\n')
                            if len(lignes) > 1:
                                for l in lignes:
                                    p2 = table.cell(x + ecart, y - nbcolonnesignorees).add_paragraph()
                                    add_run_copy(p2, r, l)
                            else:
                                add_run_copy(p2, r)
                        delete_paragraph(p)
                    # Nettoyage finale, suppression des paragraphs
                    # vides et application du style
                    for p in table.cell(x + ecart, y - nbcolonnesignorees).paragraphs:
                        if p.text == '' or p.text == ' ':
                            delete_paragraph(p)
                        else:
                            # On applique le style
                            p.style = doc.styles[textstyle]
                            # on force un alignement à GAUCHE car
                            # par défaut JUSTIFIE mais inapproprié
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # il faut au moins 1 paragraph par cellule sinon erreur
                    if len(table.cell(x + ecart, y - nbcolonnesignorees).paragraphs) < 1:
                        table.cell(x + ecart, y - nbcolonnesignorees).add_paragraph()

        # On recherche les cellules fusionnées dans le tableau Excel
        for items in sheet.merged_cells:  # Pour chaque zone fusionnée
            # on récupére les coordonnées en haut à gauche et
            # en bas à droite (rowLow,rowHigh,colLow,colHigh)
            rlo, rhi, clo, chi = items
            # on saute les entetes qui pourraient etre fusionnées
            if rlo >= nbenteteexcel:
                # on fusionne les cellules dans le tableau word en faisant
                # correspondre les numéro de ligne et de colonne
                # cf doc xlrd pour la limite superieure de mergedcell,
                # il faut -1 afin d'avoir la bonne valeur dans
                # word et on ignore la premiere colonne
                # d'excel donc -1 supplémentaire pour la colonne dans word
                table.cell(rlo + ecart, clo - 1).merge(table.cell(rhi - 1 + ecart, chi - 2))
        log.insert(END, "\ntableau du fichier Excel " + os.path.basename(
            excel) + " copié")
        return 0
    except:
        if thor["debug"]:  # si mode debug activ" # si mode debug activ"
            sys.exc_info()[0]  # On affiche l'erreur
            raise  # levée de l'erreur
        else:
            log.insert(END, "\nERROR : Erreur à la copie du tableau du fichier Excel " + os.path.basename(
                excel) + " dans le tableau " + tab[
                           "keyWord"] + ". Merci de vérifier que les formats des tableaux Word et Excel soient "
                                        "identiques")
            return 1


###################################################################
# Copy du style d'un run vers un nouveau run
# INPUT : paragraph : paragraph de destination du nouveau RUN
# INPUT : run : run source du style
# INPUT : text : optionnel : text du nouveau run,
# par défaut celui du run source est recopié


def add_run_copy(paragraph, run, text=None):
    r = paragraph.add_run(text=run.text if text is None else text,
                          style=run.style)
    r.bold = run.bold
    r.italic = run.italic
    r.underline = run.underline
    r.font.all_caps = run.font.all_caps
    r.font.bold = run.font.bold
    r.font.color.rgb = run.font.color.rgb
    r.font.color.theme_color = run.font.color.theme_color
    r.font.complex_script = run.font.complex_script
    r.font.cs_bold = run.font.cs_bold
    r.font.cs_italic = run.font.cs_italic
    r.font.double_strike = run.font.double_strike
    r.font.emboss = run.font.emboss
    r.font.hidden = run.font.hidden
    r.font.highlight_color = run.font.highlight_color
    r.font.imprint = run.font.imprint
    r.font.italic = run.font.italic
    r.font.math = run.font.math
    r.font.name = run.font.name
    r.font.no_proof = run.font.no_proof
    r.font.outline = run.font.outline
    r.font.rtl = run.font.rtl
    r.font.shadow = run.font.shadow
    r.font.size = run.font.size
    r.font.small_caps = run.font.small_caps
    r.font.snap_to_grid = run.font.snap_to_grid
    r.font.spec_vanish = run.font.spec_vanish
    r.font.strike = run.font.strike
    r.font.subscript = run.font.subscript
    r.font.superscript = run.font.superscript
    r.font.underline = run.font.underline
    r.font.web_hidden = run.font.web_hidden
    r.font.color.rgb = run.font.color.rgb
    return r


###################################################################
# Modification du style des bordures des cellules dans le tableau word
# INPUT : table : tableau word
# INPUT : color : couleur hexa de la bordure => #FFFFFF


def modifytableborders(table, width, color):
    tbl = table._tbl  # récupération de l'élément XML correspondant à la table
    for cell in tbl.iter_tcs():  # Pour cahque cellule de la table
        tcpr = cell.tcPr  # récupération de l'élément XML tcpr de la cellule
        tcborders = OxmlElement('w:tcborders')  # Création d'un element XML Borders
        # Pour chaque bordure haut, gauche, bas, droite => Ordre important
        for tag in ('w:top', 'w:left', 'w:bottom', 'w:right'):
            element = OxmlElement(tag)  # création d'un element XML correspondant à la bordure en cours
            element.set(qn('w:sz'), str(width))  # taille de la bordure
            element.set(qn('w:val'), 'single')  # bordure simple
            element.set(qn('w:color'), color)  # couleur de la bordure
            element.set(qn('w:space'), '0')  # espacement entre les bordures
            tcborders.append(element)  # ajout de l'élément à l'élément XML borders
        tcpr.append(tcborders)  # Ajout de l'élément XML BORDERS à la cellule


###################################################################
# class permettant de stocker une echelle dans un tableau


class EchelleFixe:
    def __init__(self, nom, couleur):
        self.nom = nom  # valeur de comparaison de l'echelle
        self.couleur = couleur  # couleur associée à la valeur de l'echelle


###################################################################
# Classe permetant de stocker les parametre pour une echelle calculée


class EchelleCalculee:
    def __init__(self, nom, seuil, couleur):
        self.nom = nom  # valeur de comparaison de l'echelle
        self.seuil = seuil  # seuil de déclenchement de l'echelle
        self.couleur = couleur  # couleur associée au seuil de l'echelle


###################################################################
# Classe permetant de definir une echelle


class Echelle:
    def __init__(self, echkey, methode, excel, sheet, nbentete, log, thor):
        self.methode = methode  # memorisation de la methode de l'echelle
        if methode == "fixe":
            # chargement d'une echelle fixe
            self.valeurs = load_echelle_fixe(excel, sheet, nbentete, log, thor)
        elif methode == "calculée":
            # chargement d'une echelle calculée
            self.valeurs = load_echelle_calculee(excel, sheet, nbentete, log, thor)
        else:
            if thor["debug"]:  # si mode debug activ" # si mode debug activ"
                sys.exc_info()[0]  # On affiche l'erreur
                raise  # levée de l'erreur
            else:
                log.insert(END, "\nWarning : La configuration de la \
                legende '" + echkey + "' n'est pas conforme")


###################################################################
# class permettant de récupérer la liste des scenarios stratégiques


class ScenarioStrategique:
    def __init__(self, ref, nom):
        self.ref = ref
        self.nom = nom


###################################################################
# Generationn du rapport
# INPUT : le tableau contenant tous les parametres pour la génération du script
# INPUT : log : pointeur sur a fenetre de journalisation
# INPUT : thor : configuration issue du fichier yaml


def generate_rapport(config, context, log, thor):
    log.delete('1.0', END)  # on efface la journalisation de la précédente gérération de rapport
    nberror = 0  # initalisation du nombre d'erreurs survenues lors de la génration du rapport
    echelle = {}  # initalisation du tableau des echelles
    swd = os.path.dirname(os.path.realpath(sys.argv[0]))  # repertoire d'installation du script
    try:
        if config["Rapport_input"] != '':
            doc = DocxTemplate(config["Rapport_input"])  # si le document word est renseigné
        else:
            doc = DocxTemplate(swd + "/modele/modele.docx")  # sinon on prend le modele par défaut
    except:
        if thor["debug"]:  # si mode debug activ" # si mode debug activ"
            sys.exc_info()[0]  # On affiche l'erreur
            raise  # levée de l'erreur
        else:
            log.insert(END, "\nERROR : le document word '" + config["Rapport_input"] + "' \
            ne peut pas etre ouvert")
            raise

    # chargement des legendes a partir des fichiers Excel
    for _atelierkey, atelier in thor["echelles"].items():  # niveau 2, pour chaque atelier (ici atelier echelle)
        for _titlekey, title in atelier.items():  # pour chaque regroupement d'echelles
            for echkey, ech in title.items():  # pour chaque echelle déclarée
                if config[echkey] != '':  # si le fichier excel est renseigné
                    excel = config[echkey]  # recuperation du chemin du fichier excel
                    sheet = ech["feuilleExcel"]  # recuperation du nom de la feuille Excel
                    nbentete = ech["enteteExcel"]  # nombre de ligne d'entete du fichier excel
                    # chargement de l'echelle
                    echelle[echkey] = Echelle(echkey, ech["methode"], excel, sheet, nbentete, log, thor)
                    log.insert(END, "\nechelle fixe " + echkey + " copiée")
                else:
                    log.insert(END, "\nWARNING: La légende " + echkey + " a été ignorée")
    doc.render(context)
    # Recherche des table à copier lors d'une lecture du document word.
    for x in range(0, len(doc.tables)):  # Pour chaque table
        # niveau 2 regroupement par atelier
        for _atelierkey, atelier in thor["tableaux"].items():
            # niveau 3 regroupement par sous-atelier
            for _titlekey, title in atelier.items():
                # niveau 4 - pour chaque tableau
                for tabkey, tab in title.items():
                    # si fichier Excel
                    if tab["type"] == "file" and tab["extension"] == "xls":
                        # Si la cellule [0,0] de la table correspond
                        if doc.tables[x].cell(0, 0).text == tab["keyWord"]:
                            if config[tabkey] != '':  # si le fichier excel est renseigné
                                # copie de la table du fichier Excel
                                nberror += copy_table(doc, x, tab, thor, config[tabkey], log)
                                # style des birdures
                                modifytableborders(doc.tables[x],
                                                   tab["style"]["borderWidth"],
                                                   tab["style"]["borderColor"])
                                # Si l'on a rens eigner des colonnes à styliser
                                if "colonnes" in tab["style"]:
                                    # pour chaque colonne à styliser
                                    for colKey, col in tab["style"]["colonnes"].items():
                                        # Pour chaque ligne
                                        for y in range(tab["enteteWord"],
                                                       len(doc.tables[x].rows)):
                                            # cellule de la 6éme colonne
                                            cell = doc.tables[x].cell(y, int(colKey))
                                            # si l'on a spécifier une echelle pour la colonnne
                                            if "echelle" in col:
                                                # nom de l'echelle à utiliser
                                                nom = col["echelle"]
                                                try:
                                                    # si l'echelle est configuré
                                                    if nom in echelle.keys():
                                                        ech = echelle[nom]  # echelle
                                                        # pour chaques valeurs de l'echelle
                                                        for z in range(0, len(ech.valeurs)):
                                                            # si c'est une echelle fixe
                                                            if ech.methode == "fixe":
                                                                # si le texte de la cellule correspond à la valeur de
                                                                # l'echelle
                                                                if ech.valeurs[z].nom == cell.text:
                                                                    # couleur de fond de la troisieme colonne à partir
                                                                    # de la premiere lettre du contenu de la cellule
                                                                    set_shade_cell(cell, ech.valeurs[z].couleur)
                                                            elif ech.methode == "calculée":  # si c'est une echelle
                                                                # calculée
                                                                # si le seuil correspond
                                                                if float(cell.text[0:3]) >= float(ech.valeurs[z].seuil):
                                                                    # couleur de fond de la cellule
                                                                    set_shade_cell(cell, ech.valeurs[z].couleur)
                                                except:
                                                    if thor["debug"]:  # si mode debug activ" # si mode debug activ"
                                                        sys.exc_info()[0]  # On affiche l'erreur
                                                        raise  # levée de l'erreur
                                                    else:
                                                        log.insert(END,
                                                                   "\nWARNING: La légende " + nom + " est incorrecte")
                                            if "alignment" in col:  # si l'on a preciser l'alignement du texte pour
                                                # la colonne
                                                align = col["alignment"]  # recupereation de l'alignement
                                                for p in cell.paragraphs:  # pour chaque paragraph de la cellule word
                                                    if align == "center":
                                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                    elif align == 'left':
                                                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                                    elif align == 'right':
                                                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                                    else:
                                                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                            if "backgroundColor" in col:  # si l'on a preciser la couleur de fond
                                                # pour la colonne
                                                couleur = col[
                                                    "backgroundColor"]  # récupération de la couleur au format
                                                # hexadecimal
                                                set_shade_cell(cell,
                                                               couleur)  # on applique la couleur de fond à la cellulle
                            else:
                                log.insert(END, "\nWARNING : tableau " + tab["keyWord"] + " de l’écosystème ignoré")
                    elif tab["type"] == "image":  # si image
                        # Si la cellule [0,0] de la table correspond
                        if doc.tables[x].cell(0, 0).text == tab["keyWord"]:
                            if config[tabkey] != '':  # si le fichier excel est renseigné
                                # effacement de l'ancienne illustration
                                remove_row(doc.tables[x], doc.tables[x].rows[tab["enteteWord"]])
                                doc.tables[x].add_row()  # ajout d'une ligne vierge dans le tableau word
                                doc.tables[x].cell(0, 1).paragraphs[0].add_run()  # ajout d'un run pour contenir l'image
                                # ajout de l'image
                                doc.tables[x].cell(0, 1).paragraphs[0].runs[0].add_picture(config[tabkey],
                                                                                           height=Cm(tab["height"]))
                                modifytableborders(doc.tables[x], tab["style"]["borderWidth"],
                                                   tab["style"]["borderColor"])  # style des birdures
                                # si l'on a preciser l'alignement du texte pour la colonne
                                if "alignment" in tab["style"]:
                                    align = tab["style"]["alignment"]  # recupereation de l'alignement
                                    # pour chaque paragraph de la cellule word
                                    for p in doc.tables[x].cell(0,
                                                                1).paragraphs:
                                        if align == "center":
                                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        elif align == 'left':
                                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        elif align == 'right':
                                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                        else:
                                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # sauvegarde finale
    try:
        doc.save(config["Rapport_output"])  # sauvegarde finale du rapport
    except:
        if thor["debug"]:  # si mode debug activ" # si mode debug activ"
            sys.exc_info()[0]  # On affiche l'erreur
            raise  # levée de l'erreur
        else:
            log.insert(END, "\nERROR : la sauvegarde du rapport à échouée")
            nberror = nberror + 1
    return nberror
